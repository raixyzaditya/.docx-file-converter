import fitz
import re
import sys
import json
from docx import Document
import os

def get_text_extracted(input_file):
    doc = fitz.open(input_file)
    text = ""

    for page in doc:
        text += page.get_text()
    
    word_doc = Document()
    word_doc.add_heading("Generated file :- ",level=1)
    
    for i in text.split('/n'):
        word_doc.add_paragraph(i.strip())

    save_address = os.path.splitext(input_file)[0] + ".docx"
    word_doc.save(save_address)

    print(json.dumps({
        "your_file":os.path.basename(save_address)
    }))
if __name__ == "__main__":
    file_path = sys.argv[1]
    get_text_extracted(file_path)
